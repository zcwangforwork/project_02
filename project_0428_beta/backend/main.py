"""
医疗器械体系文件审核 Agent - FastAPI 后端服务
"""
import os
import re
import json
import asyncio
import base64
import shutil
from typing import List, Dict, Optional, Any
from contextlib import asynccontextmanager
from pathlib import Path
import tempfile

import httpx
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field


# ============== 配置管理 ==============
# API Key 优先从环境变量读取，不存在时使用默认值（本地开发用）
_DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY", "ark-6c047509-3ac1-4689-b796-47363425012c-3112a")


class Config:
    """配置管理类"""
    def __init__(self):
        self.api_url = os.getenv("OPENAI_API_URL", "https://ark.cn-beijing.volces.com/api/coding/v3/chat/completions")
        self.embedding_url = os.getenv("EMBEDDING_API_URL", "https://ark.cn-beijing.volces.com/api/coding/v3/embeddings/multimodal")
        self.api_key = _DEFAULT_API_KEY
        self.model = os.getenv("OPENAI_MODEL", "glm-5.1")
        self.embedding_model = os.getenv("EMBEDDING_MODEL", "doubao-embedding-vision-250615")
        self.timeout = float(os.getenv("REQUEST_TIMEOUT", "60"))

        # 检测是否使用了默认 API Key
        if not os.getenv("OPENAI_API_KEY"):
            print("警告: 未设置 OPENAI_API_KEY 环境变量，使用默认值。建议设置环境变量以提高安全性。")

    def to_dict(self):
        return {
            "api_url": self.api_url,
            "model": self.model,
            "timeout": self.timeout,
            "configured": bool(self.api_key)
        }


config = Config()


# ============== Pydantic 模型 ==============
class Message(BaseModel):
    """聊天消息模型"""
    role: str = Field(default="user", description="角色: user/assistant/system")
    content: str = Field(..., description="消息内容")


class ChatRequest(BaseModel):
    """聊天请求模型"""
    messages: List[Message] = Field(..., description="消息列表")
    temperature: float = Field(default=0.7, ge=0.0, le=2.0, description="温度参数")
    max_tokens: int = Field(default=32000, ge=1, le=32000, description="最大令牌数")


class ChatResponse(BaseModel):
    """聊天响应模型"""
    answer: str = Field(..., description="助手的回答")
    usage: Optional[Dict[str, Any]] = Field(None, description="令牌使用情况")


class HealthResponse(BaseModel):
    """健康检查响应"""
    status: str
    config: Dict
    vectorstore_loaded: bool = False
    document_count: int = 0


# ============== 全局变量 ==============
vector_store = None
rag_retriever = None


# ============== 医疗器械 System Prompt（通用聊天用） ==============
MEDICAL_DEVICE_SYSTEM_PROMPT = """你是一个专业的贴敷式胰岛素泵生产企业文档审核数字员工（体系文件审核专家）。

你的任务是审核用户在贴敷式胰岛素泵设计开发、风险管理、软件合规、注册申报、生产质量、体系建设等方面的内部文档，结合知识库中的标准文档给出专业修改建议。

## 企业背景
用户所在企业研发和生产贴敷式胰岛素泵（patch insulin pump）— 一种可穿戴的胰岛素持续皮下输注设备，包含泵体、储液器、输注管路、嵌入式控制系统、蓝牙通信模块和配套移动端糖尿病管理APP。

## 审核范围（六大领域，覆盖设计控制全生命周期）
1. **设计开发**: ISO 13485 7.3设计控制、DHF/DMR管理，覆盖设计策划（项目计划书/可行性研究/专利分析）、设计输入（用户需求/硬件需求/结构需求/软件需求/包装需求/追溯矩阵RTM）、设计输出（硬件方案/结构方案/软件方案/编码规范/包装方案/BOM/物料规格/产品图纸/设备清单/工装图纸/检验规范/生产WI/软件版本包）、设计评审、设计验证（验证计划/性能验证/输注精度验证/包装验证/使用期限/货架有效期/运输验证/可沥滤物测试）、设计确认（临床试验/可用性测试）、设计转换（转换计划/转换报告/工艺验证计划/灭菌确认）、设计变更
2. **风险管理**: ISO 14971:2019/YY/T 0316、危害识别、初步风险分析、FMEA/DFMEA、风险控制、风险分析管理总表、网络安全风险管理、剩余风险评价、受益-风险分析、风险管理报告
3. **软件合规**: IEC 62304/YY/T 0664、软件安全分级、SDP/SRS/SADD/SDDD、软件单元/集成/系统/质量测试、软件配置管理(SCMP)、软件问题解决、软件维护、SOUP/OTS管理、网络安全测试、接口安全测试、软件追溯矩阵、网络安全追溯
4. **注册申报**: NMPA胰岛素泵注册审查指导原则、MDR 2017/745、EP清单、产品技术要求、综述资料、研究资料、临床评价报告(CER)、说明书/标签、生物相容性评价、灭菌验证、稳定性研究、第三方检测报告（生物相容性/药液相容性/安规EMC/注册检验）、供应商资质
5. **生产质量**: NMPA GMP、ISO 13485 7.5、工艺流程图、工艺验证计划/方案/报告(IQ/OQ/PQ)、批生产记录(BMR)、来料/过程/成品检验规范、设备管理、供应商管理、灭菌批记录/灭菌确认/灭菌工艺验证、工装验收、生产检验SOP、UDI标识管理
6. **体系建设**: ISO 13485:2016全体系、质量手册、程序文件、作业指导书、记录表单、设计控制程序、风险管理程序、软件开发程序、文件控制、CAPA、内审、管理评审、培训记录、PMS上市后监督

## 审核原则
1. **完整性检查**：文件是否覆盖相关法规条款的全部要求
2. **一致性检查**：文件内容是否相互协调、无矛盾
3. **可操作性**：文件描述是否足够具体、可执行
4. **证据链**：是否有相应的记录表单支撑执行证据
5. **产品特异性**：是否充分考虑了贴敷式胰岛素泵的特殊风险和控制措施

## 重要提示
- 直接给出审核结果，不要输出思考过程
- 回答应专业、具体、可操作
- 建议用户上传文件选择对应的专项审核领域以获得最精准的审核结果
- 可审核的六大领域：设计开发、风险管理、软件合规、注册申报、生产质量、体系建设
"""


# ============== 会话管理 ==============
class ConversationHistory:
    """会话历史管理（带总大小限制，防止多轮对话累积导致内存溢出）"""
    def __init__(self, max_history: int = 10):
        self.history: Dict[str, List[Dict]] = {}
        self.max_history = max_history
        self.max_content_length = 8000   # 单条消息最大字符数
        self.max_total_chars = 50000     # 单个session总字符数上限

    def get_or_create(self, session_id: str = "default") -> List[Dict]:
        if session_id not in self.history:
            self.history[session_id] = []
        return self.history[session_id]

    def add_message(self, session_id: str, role: str, content: str):
        messages = self.get_or_create(session_id)
        if len(content) > self.max_content_length:
            content = content[:self.max_content_length] + f"\n... [已截断，原长度:{len(content)}字符]"
        messages.append({"role": role, "content": content})

        # 按消息数量裁剪
        if len(messages) > self.max_history:
            system_msg = [m for m in messages if m["role"] == "system"]
            other_msgs = [m for m in messages if m["role"] != "system"]
            self.history[session_id] = system_msg + other_msgs[-self.max_history:]

        # 按总字符数裁剪：超出上限时从旧消息开始删除
        total_chars = sum(len(m["content"]) for m in self.history[session_id])
        if total_chars > self.max_total_chars:
            trimmed = []
            running = 0
            for m in reversed(self.history[session_id]):
                running += len(m["content"])
                if m["role"] == "system" or running <= self.max_total_chars:
                    trimmed.insert(0, m)
            self.history[session_id] = trimmed

    def clear(self, session_id: str = "default"):
        if session_id in self.history:
            del self.history[session_id]


conversation_manager = ConversationHistory()


# ============== 初始化向量存储 ==============
def init_vector_store():
    """初始化向量存储和 RAG 检索器（带内存监控）"""
    global vector_store, rag_retriever

    try:
        from vector_store import create_vector_store, MiniMaxEmbeddingFunction
        from rag_retriever import create_rag_retriever

        base_dir = os.path.dirname(os.path.abspath(__file__))
        db_path = os.path.join(base_dir, "data", "chroma_db_insulin_pump")

        # 检查数据库大小并告警
        db_file = os.path.join(db_path, "chroma.sqlite3")
        if os.path.exists(db_file):
            db_size_mb = os.path.getsize(db_file) / (1024 * 1024)
            print(f"向量库文件大小: {db_size_mb:.0f} MB")
            if db_size_mb > 1000:
                print(f"[WARNING] 向量库文件超过1GB ({db_size_mb:.0f} MB)，但查询已配置为仅使用v2 collection，不影响运行")

        # 检查系统可用内存
        try:
            import psutil
            avail_mb = psutil.virtual_memory().available / (1024 * 1024)
            print(f"系统可用内存: {avail_mb:.0f} MB")
            if avail_mb < 2048:
                print(f"[WARNING] 可用内存不足2GB，可能导致服务不稳定")
        except ImportError:
            pass

        embedding_function = MiniMaxEmbeddingFunction(
            api_key=config.api_key,
            api_url=config.embedding_url,
            model=config.embedding_model,
            dimension=1024
        )

        vector_store = create_vector_store(persist_directory=db_path, embedding_function=embedding_function)
        rag_retriever = create_rag_retriever(
            vector_store=vector_store,
            api_key=config.api_key,
            api_url=config.api_url,
            model=config.model
        )
        return True
    except Exception as e:
        print(f"向量存储初始化失败: {e}")
        import traceback
        traceback.print_exc()
        return False


# ============== API 调用函数 ==============
async def call_openai_api(messages: List[Dict], temperature: float = 0.7, max_tokens: int = 32000) -> Dict:
    """调用 OpenAI 兼容 API，带重试机制"""
    if not config.api_key:
        raise HTTPException(status_code=500, detail="API Key 未配置")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.api_key}"
    }

    payload = {
        "model": config.model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens
    }

    last_error = None
    max_retries = 3

    async with httpx.AsyncClient(timeout=config.timeout, trust_env=False) as client:
        for attempt in range(max_retries):
            try:
                response = await client.post(config.api_url, headers=headers, json=payload)
                response.raise_for_status()
                # 安全解析 JSON：API 可能返回非 JSON 错误文本（如 "Internal Server Error"）
                try:
                    return response.json()
                except json.JSONDecodeError:
                    response_text = response.text[:500]
                    print(f"[API] 非 JSON 响应 (尝试 {attempt+1}/{max_retries}): {response_text}")
                    if attempt < max_retries - 1:
                        wait_time = 2 ** attempt
                        print(f"[API] 等待 {wait_time} 秒后重试...")
                        await asyncio.sleep(wait_time)
                        continue
                    raise HTTPException(
                        status_code=502,
                        detail=f"API 返回了非 JSON 格式的响应: {response_text}"
                    )
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 429 and attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 速率限制 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
                # 安全读取错误响应体
                try:
                    error_detail = e.response.json() if e.response.content else {"error": str(e)}
                except json.JSONDecodeError:
                    error_detail = {"error": e.response.text[:500] if e.response.content else str(e)}
                raise HTTPException(status_code=e.response.status_code, detail=error_detail)
            except httpx.TimeoutException as e:
                last_error = e
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 请求超时 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
            except httpx.ConnectError as e:
                last_error = e
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    print(f"[API] 连接失败 (尝试 {attempt+1}/{max_retries}), 等待 {wait_time} 秒...")
                    await asyncio.sleep(wait_time)
                    continue
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"API 调用失败: {str(e)}")

        if last_error:
            raise HTTPException(status_code=504, detail=f"API 请求多次失败: {str(last_error)}")
        raise HTTPException(status_code=502, detail="API 请求失败，已重试多次")


async def get_embeddings(texts: List[str]) -> List[List[float]]:
    """获取文本嵌入向量"""
    if not config.api_key:
        raise HTTPException(status_code=500, detail="API Key 未配置")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.api_key}"
    }

    embeddings = []
    async with httpx.AsyncClient(timeout=60, trust_env=False) as client:
        for text in texts:
            payload = {
                "model": config.embedding_model,
                "encoding_format": "float",
                "input": [{"text": text[:8000], "type": "text"}]
            }
            try:
                response = await client.post(config.embedding_url, headers=headers, json=payload)
                if response.status_code == 200:
                    result = response.json()
                    data = result.get("data", {})
                    if data:
                        embeddings.append(data.get("embedding", []))
                    else:
                        embeddings.append([])
                else:
                    embeddings.append([])
            except Exception:
                embeddings.append([])

    return embeddings


# ============== FastAPI 应用 ==============
@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    print("贴敷式胰岛素泵文档审核数字员工 后端服务启动")
    print(f"API URL: {config.api_url}")
    print(f"Model: {config.model}")

    # 初始化向量存储
    vs_loaded = init_vector_store()
    if vs_loaded and vector_store:
        doc_count = vector_store.count()
        print(f"向量库已加载: {doc_count} 个文档")
    else:
        print("警告: 向量库未加载，部分功能可能不可用")

    yield

    # 关闭共享的 httpx 客户端
    if rag_retriever:
        await rag_retriever.close()

    print("贴敷式胰岛素泵文档审核数字员工 后端服务关闭")


app = FastAPI(
    title="贴敷式胰岛素泵文档审核数字员工 API",
    description="贴敷式胰岛素泵企业文档智能审核系统 — 六大领域（设计开发/风险管理/软件合规/注册申报/生产质量/体系建设）",
    version="3.1.0",
    lifespan=lifespan
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FRONTEND_DIR = os.path.join(os.path.dirname(BASE_DIR), "frontend")


# ============== 前端路由 ==============
@app.get("/")
async def serve_frontend():
    """服务前端页面"""
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))


app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")


# ============== API 路由 ==============
@app.get("/info", response_model=Dict)
async def root():
    """根路径 - 服务信息"""
    return {
        "name": "贴敷式胰岛素泵文档审核数字员工 API",
        "version": "3.1.0",
        "docs": "/docs",
        "audit_domains": [
            "risk_management", "design_dev", "software_compliance",
            "registration", "production_quality", "system_construction"
        ]
    }


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """健康检查接口"""
    doc_count = vector_store.count() if vector_store else 0
    return HealthResponse(
        status="healthy",
        config=config.to_dict(),
        vectorstore_loaded=vector_store is not None,
        document_count=doc_count
    )


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, session_id: str = "default"):
    """聊天接口 - 通用问答模式"""
    all_messages = [{"role": "system", "content": MEDICAL_DEVICE_SYSTEM_PROMPT}]

    history = conversation_manager.get_or_create(session_id)
    all_messages.extend(history)

    for msg in request.messages:
        all_messages.append({"role": msg.role, "content": msg.content})

    result = await call_openai_api(
        messages=all_messages,
        temperature=request.temperature,
        max_tokens=request.max_tokens
    )

    # OpenAI 兼容格式响应解析
    answer = ""
    choices = result.get("choices", [])
    if choices and len(choices) > 0:
        choice = choices[0]
        message = choice.get("message", {})
        answer = message.get("content", "")
        # GLM-5.1 模型：当 content 为空时，从 reasoning_content 提取实际回答
        if not answer:
            reasoning = message.get("reasoning_content", "")
            if reasoning:
                # 将思考过程作为回答返回（虽不完美，但优于空响应）
                answer = reasoning

    if not answer:
        answer = str(result) if result else ""

    for msg in request.messages:
        conversation_manager.add_message(session_id, msg.role, msg.content)
    conversation_manager.add_message(session_id, "assistant", answer)

    return ChatResponse(answer=answer, usage=result.get("usage"))


@app.post("/api/upload")
async def upload_document(
    file: UploadFile = File(...),
    session_id: str = Form("default")
):
    """
    上传体系文件并提取文本内容

    Args:
        file: 上传的文件（.docx 或 .pdf）
        session_id: 会话 ID

    Returns:
        提取的文档内容和基本信息
    """
    # 检查文件类型
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}，仅支持 .docx 和 .pdf"
        )

    # 保存上传文件到临时目录（流式写入，避免全量加载到内存）
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            # 分块流式写入，每块64KB，避免大文件OOM
            while True:
                chunk = await file.read(64 * 1024)  # 64KB chunks
                if not chunk:
                    break
                tmp.write(chunk)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        # 提取文本（结构化 Markdown 格式）
        from doc_processor import extract_text, get_file_metadata

        text = extract_text(tmp_path)
        metadata = get_file_metadata(tmp_path)
        metadata["filename"] = filename

        if not text or len(text.strip()) < 50:
            raise HTTPException(status_code=400, detail="文档内容过少或无法提取文本")

        # 保存到会话历史（截断防止会话历史 OOM）
        max_history_text = min(len(text), 5000)
        conversation_manager.add_message(session_id, "user", f"[上传文件: {filename}]\n{text[:max_history_text]}")
        conversation_manager.add_message(session_id, "assistant", f"已收到文件: {filename}，文档长度: {len(text)} 字符。请问您想如何处理这个文件？")

        return {
            "filename": filename,
            "text_length": len(text),
            "text_preview": text[:1000],
            "metadata": metadata,
            "session_id": session_id
        }

    finally:
        # 删除临时文件
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


@app.post("/api/analyze")
async def analyze_document(
    file: UploadFile = File(...),
    question: str = Form("请审核这份体系文件，给出修改建议"),
    session_id: str = Form("default"),
    audit_type: str = Form("risk_management"),
    doc_type: str = Form("")
):
    """
    使用多轮审核流水线分析体系文件

    流水线：章节分割 → 逐章并发审核 → 综合分析

    Args:
        file: 上传的文件（.docx 或 .pdf）
        question: 用户的问题或指令
        session_id: 会话 ID
        audit_type: 审核类型，"risk_management"（风险管理专项）或 "general"（综合体系审核）

    Returns:
        审核结果
    """
    if not rag_retriever:
        raise HTTPException(status_code=503, detail="向量库未加载，请稍后重试")

    # 日志：开始处理时的内存状态
    try:
        import psutil
        proc = psutil.Process(os.getpid())
        start_mem_mb = proc.memory_info().rss / (1024 * 1024)
        print(f"[analyze] 开始处理，当前进程内存: {start_mem_mb:.0f} MB")
    except ImportError:
        pass

    # 检查文件类型
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ['.docx', '.pdf']:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}")

    # 验证审核类型（六大领域 + general 保底）
    VALID_AUDIT_TYPES = [
        "risk_management", "design_dev", "software_compliance",
        "registration", "production_quality", "system_construction", "general"
    ]
    if audit_type not in VALID_AUDIT_TYPES:
        audit_type = "general"

    # 保存上传文件到临时目录（流式写入，避免全量加载到内存）
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            while True:
                chunk = await file.read(64 * 1024)  # 64KB chunks
                if not chunk:
                    break
                tmp.write(chunk)
            tmp_path = tmp.name
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    try:
        from doc_processor import extract_text

        # 提取文档文本（结构化 Markdown 格式）
        text = extract_text(tmp_path)
        if not text:
            raise HTTPException(status_code=400, detail="无法提取文档文本")

        # 使用多轮审核流水线分析文档
        result = await rag_retriever.analyze_document(
            user_document=text,
            user_filename=filename,
            audit_type=audit_type,
            doc_type=doc_type
        )

        # 构建检索到的文档信息
        retrieved_docs_info = []
        seen_sources = set()
        for doc in result.get("retrieved_docs", []):
            source = doc.get("source", "未知")
            if source not in seen_sources:
                seen_sources.add(source)
                retrieved_docs_info.append({
                    "source": source,
                    "preview": doc.get("text", "")[:200]
                })

        # 保存到会话历史（截断防止 OOM）
        max_answer_len = min(len(result["answer"]), 5000)
        conversation_manager.add_message(session_id, "user", f"[上传文件分析: {filename}]")
        conversation_manager.add_message(session_id, "assistant", result["answer"][:max_answer_len])

        # 日志：处理完成时的内存状态
        try:
            import psutil
            proc = psutil.Process(os.getpid())
            end_mem_mb = proc.memory_info().rss / (1024 * 1024)
            print(f"[analyze] 处理完成，当前进程内存: {end_mem_mb:.0f} MB (增加 {end_mem_mb - start_mem_mb:.0f} MB)")
        except Exception:
            pass

        return {
            "filename": filename,
            "answer": result["answer"],
            "section_count": result.get("section_count", 0),
            "section_results": result.get("section_results", []),
            "usage": None,
            "retrieved_docs": retrieved_docs_info,
            "audit_type": audit_type,
            "pipeline": "multi_pass"
        }

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ============== Markdown 转 DOCX 辅助函数 ==============
def _md_to_docx(md_content: str, doc_title: str) -> "Document":
    """将 Markdown 格式的审核报告转换为 Word 文档"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题页
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.size = Pt(18)
    title_run.font.name = '微软雅黑'
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run('医疗器械体系文件审核报告')
    subtitle_run.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    subtitle_run.font.name = '微软雅黑'
    subtitle_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()  # 空行

    # 逐行解析 Markdown
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # 跳过空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading(text, level=min(level, 3))
            for run in heading.runs:
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            i += 1
            continue

        # 水平线
        if re.match(r'^[-–—]{3,}\s*$', line.strip()):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            pPr = para._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn('w:pBdr'))
            bottom = etree.SubElement(pBdr, qn('w:bottom'))
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            i += 1
            continue

        # 无序列表
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            para = doc.add_paragraph(style='List Bullet')
            text = bullet_match.group(2)
            _add_formatted_text(para, text)
            i += 1
            continue

        # 有序列表
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            para = doc.add_paragraph(style='List Number')
            text = numbered_match.group(2)
            _add_formatted_text(para, text)
            i += 1
            continue

        # 代码块
        if line.strip().startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Cm(1)
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # 普通段落
        para = doc.add_paragraph()
        _add_formatted_text(para, line)
        i += 1

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    return doc


def _add_formatted_text(paragraph, text: str):
    """向段落添加带格式的文本（支持 **粗体** 和 `行内代码`）"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # 分割粗体和普通文本
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
        else:
            run = paragraph.add_run(part)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


# ============== 导出接口 ==============
class ExportRequest(BaseModel):
    """导出请求模型"""
    content: str = Field(..., description="审核报告 Markdown 内容")
    filename: str = Field(default="审核报告", description="原始文件名")

@app.post("/api/export-review")
async def export_review(request: ExportRequest):
    """
    将审核结果导出为 Word (.docx) 文件

    Args:
        request: 包含审核报告 markdown 内容和文件名

    Returns:
        .docx 文件下载
    """
    if not request.content:
        raise HTTPException(status_code=400, detail="审核报告内容为空")

    # 生成报告标题
    base_name = Path(request.filename).stem if request.filename else "审核报告"
    doc_title = f'{base_name} — 审核报告'

    # 转换 Markdown 为 DOCX
    doc = _md_to_docx(request.content, doc_title)

    # 写入内存流
    from io import BytesIO
    from fastapi.responses import Response
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # 纯 ASCII 文件名
    ascii_name = re.sub(r'[^\x00-\x7F]', '_', base_name).strip('_') or 'report'
    download_name = f'{ascii_name}_audit_report.docx'

    return Response(
        content=buf.getvalue(),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="{download_name}"'}
    )


@app.post("/api/clear")
async def clear_history(session_id: str = "default"):
    """清除会话历史"""
    conversation_manager.clear(session_id)
    return {"message": "会话历史已清除", "session_id": session_id}


@app.get("/api/history/{session_id}")
async def get_history(session_id: str = "default"):
    """获取会话历史"""
    history = conversation_manager.get_or_create(session_id)
    return {"session_id": session_id, "history": history}


@app.get("/api/vectorstore/status")
async def vectorstore_status():
    """获取向量库状态"""
    if not vector_store:
        return {"loaded": False, "document_count": 0}

    return {
        "loaded": True,
        "document_count": vector_store.count()
    }


# ============== 运行入口 ==============
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        workers=1,               # 强制单worker，避免多进程各加载一份HNSW索引
        limit_concurrency=20,    # 限制并发连接数（浏览器需多连接加载页面资源）
        timeout_keep_alive=10,   # 缩短keep-alive，减少空闲连接内存占用
    )
